import os
from langchain_openai import ChatOpenAI
from docx import Document
from dotenv import load_dotenv
from app.models.database import supabase
from app.config import settings
from app.models.schemas import DocumentType, AreaOfLaw
from typing import Optional

# Load environment variables
load_dotenv()
openai_api_key = os.getenv("OPENAI_API_KEY")

if not openai_api_key:
    raise ValueError("Missing OPENAI_API_KEY in .env file")

# Initialize GPT-4.5 Turbo
llm = ChatOpenAI(model_name="gpt-4.1", temperature=0.5)

def fetch_template_from_supabase(template_path: str) -> str:
    """
    Fetch a template file from Supabase storage.

    Parameters:
    - template_path (str): Path to the template in Supabase storage.

    Returns:
    - str: Local path to the downloaded template file.
    """
    try:
        response = supabase.storage.from_("legal-templates").download(template_path)
        if response.get("error"):
            raise ValueError(f"Error downloading template: {response['error']['message']}")

        local_path = f"/tmp/{os.path.basename(template_path)}"
        with open(local_path, "wb") as file:
            file.write(response["data"])
        return local_path
    except Exception as e:
        raise RuntimeError(f"Error fetching template: {str(e)}")

def edit_template_with_ai(template_path: str, user_command: str, user_id: str, title: str) -> str:
    """
    Edit a template using AI based on user commands and save it to Supabase.

    Parameters:
    - template_path (str): Local path to the template file.
    - user_command (str): User's command to modify the template.
    - user_id (str): ID of the user.
    - title (str): Title of the document.

    Returns:
    - str: Path to the updated template in storage.
    """
    try:
        # Load the template
        doc = Document(template_path)
        template_text = "\n".join([para.text for para in doc.paragraphs])

        # Create AI prompt
        prompt = f"Here is a legal document template:\n{template_text}\n\nUser command: {user_command}\n\nModify the document accordingly and return the updated content."

        # Get AI response
        updated_content = llm.predict(prompt)

        # Update the document with AI response
        updated_doc = Document()
        for line in updated_content.split("\n"):
            updated_doc.add_paragraph(line)

        # Save the updated document
        updated_path = f"/tmp/updated_{os.path.basename(template_path)}"
        updated_doc.save(updated_path)

        # Upload to Supabase storage
        with open(updated_path, "rb") as file:
            storage_path = f"updated_{os.path.basename(template_path)}"
            supabase.storage.from_("legal-templates").upload(
                storage_path,
                file.read(),
                {"content-type": "application/vnd.openxmlformats-officedocument.wordprocessingml.document"}
            )

        # Update template record in database
        supabase.from_("templates").update({
            "content": updated_content,
            "file_path": storage_path
        }).eq("file_path", template_path).execute()

        return storage_path
    except Exception as e:
        raise RuntimeError(f"Error editing template: {str(e)}")

def get_template_content_by_criteria(
    document_type: DocumentType,
    area_of_law: AreaOfLaw,
    jurisdiction: Optional[str] = None
) -> Optional[str]:
    """
    Fetches a document template from Supabase based on document type, area of law, and optional jurisdiction.

    Parameters:
    - document_type (DocumentType): The type of document.
    - area_of_law (AreaOfLaw): The area of law.
    - jurisdiction (Optional[str]): The jurisdiction for the template (e.g., 'Florida').

    Returns:
    - Optional[str]: The content of the template if found, otherwise None.
    """
    try:
        query = supabase.from_("document_templates").select("template_text").eq("document_type", document_type.value).eq("area_of_law", area_of_law.value)
        
        if jurisdiction:
            query = query.eq("jurisdiction", jurisdiction)
        
        # Order by created_at descending to get the latest template if multiple match
        response = query.order("created_at", desc=True).limit(1).execute()

        if response.data and len(response.data) > 0:
            return response.data[0]["template_text"]
        else:
            return None
    except Exception as e:
        raise RuntimeError(f"Error fetching template by criteria: {str(e)}")
