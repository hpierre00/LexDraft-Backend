from fastapi import APIRouter, HTTPException, UploadFile, File, Depends, Form, Query
from fastapi.security import HTTPAuthorizationCredentials
from typing import List, Optional, Dict, Any
from app.models.database import supabase
from app.utils.auth_utils import get_current_user, security
from docx import Document
import io
import logging
import traceback
import os
from fastapi.responses import Response
import json

# Configure logging
logger = logging.getLogger(__name__)

router = APIRouter()

# Helper to validate state and document type
def validate_state_and_document_type(state: str, document_type: str):
    # Validate state
    state_check = supabase.from_("states").select("state_id").eq("state_name", state).execute()
    if not state_check.data:
        raise HTTPException(status_code=400, detail="Invalid state provided")

    # Validate document type (case-insensitive match)
    doc_type_check = supabase.from_("document_types").select("document_type_id").ilike("document_type_name", document_type).execute()
    if not doc_type_check.data:
        raise HTTPException(status_code=400, detail="Invalid document type provided")

# Create Template
@router.post("/create", tags=["Templates"])
async def create_template(
    state: str = Form(...),
    document_type: str = Form(...),
    file: UploadFile = Form(...),
    user: dict = Depends(get_current_user),
    credentials: HTTPAuthorizationCredentials = Depends(security)
) -> Dict[str, Any]:
    """
    Create a new template.
    
    Example form data:
    ```
    state: "California"
    document_type: "Non-Disclosure Agreement"
    file: [template.docx]
    ```
    
    Returns:
    ```json
    {
        "message": "Template created successfully",
        "file_path": "legal-templates/California/Non-Disclosure Agreement/template.docx"
    }
    ```
    """
    try:
        logger.info(f"Creating template for user {user['id']}: {file.filename}")
        
        if not file.filename.endswith('.docx'):
            raise HTTPException(
                status_code=400,
                detail="Only .docx files are supported"
            )

        # Validate state and document type
        try:
            state_response = supabase.from_("states").select("state_id").eq("state_name", state).single().execute()
            if not state_response.data:
                raise HTTPException(
                    status_code=400,
                    detail=f"Invalid state: {state}"
                )
            state_id = state_response.data["state_id"]
            logger.info(f"Validated state: {state} (ID: {state_id})")
        except Exception as state_error:
            logger.error(f"State validation failed: {str(state_error)}")
            raise HTTPException(
                status_code=400,
                detail=f"Error validating state: {str(state_error)}"
            )

        try:
            doc_type_response = supabase.from_("document_types").select("document_type_id").eq("document_type_name", document_type).single().execute()
            if not doc_type_response.data:
                raise HTTPException(
                    status_code=400,
                    detail=f"Invalid document type: {document_type}"
                )
            doc_type_id = doc_type_response.data["document_type_id"]
            logger.info(f"Validated document type: {document_type} (ID: {doc_type_id})")
        except Exception as doc_type_error:
            logger.error(f"Document type validation failed: {str(doc_type_error)}")
            raise HTTPException(
                status_code=400,
                detail=f"Error validating document type: {str(doc_type_error)}"
            )

        # Create storage path
        storage_path = f"legal-templates/{state}/{document_type}/{file.filename}"
        
        # Extract content from DOCX
        try:
            content = ""
            with open(f"/tmp/{file.filename}", "wb") as f:
                f.write(await file.read())
            doc = Document(f"/tmp/{file.filename}")
            content = "\n".join([para.text for para in doc.paragraphs])
            logger.info(f"Extracted content from {file.filename}")
        except Exception as content_error:
            logger.error(f"Content extraction failed: {str(content_error)}")
            raise HTTPException(
                status_code=400,
                detail=f"Error extracting content from file: {str(content_error)}"
            )

        # Upload to Supabase storage
        try:
            with open(f"/tmp/{file.filename}", "rb") as f:
                supabase.storage.from_("legal-templates").upload(
                    storage_path,
                    f.read(),
                    {"content-type": "application/vnd.openxmlformats-officedocument.wordprocessingml.document"}
                )
            logger.info(f"File uploaded to storage: {storage_path}")
        except Exception as upload_error:
            logger.error(f"File upload failed: {str(upload_error)}")
            raise HTTPException(
                status_code=500,
                detail=f"Error uploading file: {str(upload_error)}"
            )

        # Insert into database
        try:
            template_data = {
                "state_id": state_id,
                "document_type_id": doc_type_id,
                "template_name": file.filename,
                "file_path": storage_path,
                "content": content
            }
            response = supabase.from_("templates").insert(template_data).execute()
            logger.info(f"Template registered in database: {response.data[0]['id']}")
        except Exception as db_error:
            logger.error(f"Database insert failed: {str(db_error)}")
            # Try to clean up uploaded file
            try:
                supabase.storage.from_("legal-templates").remove([storage_path])
            except:
                pass
            raise HTTPException(
                status_code=500,
                detail=f"Error registering template: {str(db_error)}"
            )

        return {
            "message": "Template created successfully",
            "file_path": storage_path
        }
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Unexpected error in create_template: {str(e)}\n{traceback.format_exc()}")
        raise HTTPException(
            status_code=500,
            detail=f"Internal server error: {str(e)}"
        )

# List Templates with Filters
@router.get("/list", tags=["Templates"])
async def list_templates(
    state_id: Optional[str] = Query(None, description="Filter by state ID"),
    document_type_id: Optional[str] = Query(None, description="Filter by document type ID"),
    search: Optional[str] = Query(None, description="Search in template name or content"),
    user: dict = Depends(get_current_user)
) -> Dict[str, List[Dict[str, Any]]]:
    """
    List templates with optional filters.
    
    Query parameters:
    - state_id: Filter by state ID
    - document_type_id: Filter by document type ID
    - search: Search in template name or content
    """
    try:
        logger.info(f"Listing templates for user {user['id']} with filters: state_id={state_id}, doc_type_id={document_type_id}, search={search}")
        
        query = supabase.from_("templates").select("""
            *,
            states:state_id(state_id, state_name),
            document_types:document_type_id(document_type_id, document_type_name)
        """)

        if state_id:
            query = query.eq("state_id", state_id)
        if document_type_id:
            query = query.eq("document_type_id", document_type_id)
        if search:
            query = query.or_(f"template_name.ilike.%{search}%,content.ilike.%{search}%")

        try:
            response = query.execute()
            logger.info(f"Found {len(response.data)} templates")
            return {"templates": response.data}
        except Exception as query_error:
            logger.error(f"Query failed: {str(query_error)}")
            raise HTTPException(
                status_code=500,
                detail=f"Error fetching templates: {str(query_error)}"
            )
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Unexpected error in list_templates: {str(e)}\n{traceback.format_exc()}")
        raise HTTPException(
            status_code=500,
            detail=f"Internal server error: {str(e)}"
        )

# Get Template by ID with Full Details
@router.get("/{template_id}", tags=["Templates"])
async def get_template(
    template_id: str,
    user: dict = Depends(get_current_user)
) -> Dict[str, Any]:
    """
    Get a specific template by ID with full details including state and document type info.
    """
    try:
        logger.info(f"Getting template {template_id} for user {user['id']}")
        
        try:
            response = supabase.from_("templates").select("""
                *,
                states:state_id(state_id, state_name),
                document_types:document_type_id(document_type_id, document_type_name)
            """).eq("id", template_id).single().execute()
            
            if not response.data:
                raise HTTPException(
                    status_code=404,
                    detail=f"Template not found: {template_id}"
                )
            
            logger.info(f"Found template: {template_id}")
            return response.data
        except Exception as query_error:
            logger.error(f"Query failed: {str(query_error)}")
            raise HTTPException(
                status_code=500,
                detail=f"Error fetching template: {str(query_error)}"
            )
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Unexpected error in get_template: {str(e)}\n{traceback.format_exc()}")
        raise HTTPException(
            status_code=500,
            detail=f"Internal server error: {str(e)}"
        )

# Get Templates by State
@router.get("/state/{state_id}", tags=["Templates"])
async def get_templates_by_state(
    state_id: str,
    user: dict = Depends(get_current_user)
) -> Dict[str, List[Dict[str, Any]]]:
    """
    Get all templates for a specific state.
    """
    try:
        logger.info(f"Getting templates for state {state_id} for user {user['id']}")
        
        try:
            response = supabase.from_("templates").select("""
                *,
                document_types:document_type_id(document_type_id, document_type_name)
            """).eq("state_id", state_id).execute()
            
            logger.info(f"Found {len(response.data)} templates for state {state_id}")
            return {"templates": response.data}
        except Exception as query_error:
            logger.error(f"Query failed: {str(query_error)}")
            raise HTTPException(
                status_code=500,
                detail=f"Error fetching templates: {str(query_error)}"
            )
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Unexpected error in get_templates_by_state: {str(e)}\n{traceback.format_exc()}")
        raise HTTPException(
            status_code=500,
            detail=f"Internal server error: {str(e)}"
        )

# Get Templates by Document Type
@router.get("/type/{document_type_id}", tags=["Templates"])
async def get_templates_by_document_type(
    document_type_id: str,
    user: dict = Depends(get_current_user)
) -> Dict[str, List[Dict[str, Any]]]:
    """
    Get all templates of a specific document type.
    """
    try:
        logger.info(f"Getting templates for document type {document_type_id} for user {user['id']}")
        
        try:
            response = supabase.from_("templates").select("""
                *,
                states:state_id(state_id, state_name)
            """).eq("document_type_id", document_type_id).execute()
            
            logger.info(f"Found {len(response.data)} templates for document type {document_type_id}")
            return {"templates": response.data}
        except Exception as query_error:
            logger.error(f"Query failed: {str(query_error)}")
            raise HTTPException(
                status_code=500,
                detail=f"Error fetching templates: {str(query_error)}"
            )
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Unexpected error in get_templates_by_document_type: {str(e)}\n{traceback.format_exc()}")
        raise HTTPException(
            status_code=500,
            detail=f"Internal server error: {str(e)}"
        )

# Search Templates
@router.get("/search", tags=["Templates"])
async def search_templates(
    query: str = Query(...),
    user: dict = Depends(get_current_user)
) -> Dict[str, List[Dict[str, Any]]]:
    """
    Search templates by name or content.
    """
    try:
        logger.info(f"Searching templates for user {user['id']} with query: {query}")
        
        try:
            response = supabase.from_("templates").select("""
                *,
                states:state_id(state_id, state_name),
                document_types:document_type_id(document_type_id, document_type_name)
            """).or_(f"template_name.ilike.%{query}%,content.ilike.%{query}%").execute()
            
            logger.info(f"Found {len(response.data)} templates matching query")
            return {"templates": response.data}
        except Exception as query_error:
            logger.error(f"Query failed: {str(query_error)}")
            raise HTTPException(
                status_code=500,
                detail=f"Error searching templates: {str(query_error)}"
            )
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Unexpected error in search_templates: {str(e)}\n{traceback.format_exc()}")
        raise HTTPException(
            status_code=500,
            detail=f"Internal server error: {str(e)}"
        )

# Delete Template
@router.delete("/{template_id}", tags=["Templates"])
async def delete_template(
    template_id: str,
    user: dict = Depends(get_current_user)
) -> Dict[str, str]:
    """
    Delete a template by ID.
    """
    try:
        logger.info(f"Deleting template {template_id} for user {user['id']}")
        
        # Get template info first
        try:
            template = supabase.from_("templates").select("file_path").eq("id", template_id).single().execute()
            if not template.data:
                raise HTTPException(
                    status_code=404,
                    detail=f"Template not found: {template_id}"
                )
            file_path = template.data["file_path"]
        except Exception as query_error:
            logger.error(f"Template fetch failed: {str(query_error)}")
            raise HTTPException(
                status_code=500,
                detail=f"Error fetching template: {str(query_error)}"
            )

        # Delete from storage
        try:
            supabase.storage.from_("legal-templates").remove([file_path])
            logger.info(f"Deleted file from storage: {file_path}")
        except Exception as storage_error:
            logger.error(f"Storage deletion failed: {str(storage_error)}")
            # Continue with database deletion even if storage deletion fails

        # Delete from database
        try:
            supabase.from_("templates").delete().eq("id", template_id).execute()
            logger.info(f"Deleted template from database: {template_id}")
        except Exception as db_error:
            logger.error(f"Database deletion failed: {str(db_error)}")
            raise HTTPException(
                status_code=500,
                detail=f"Error deleting template: {str(db_error)}"
            )

        return {"message": "Template deleted successfully"}
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Unexpected error in delete_template: {str(e)}\n{traceback.format_exc()}")
        raise HTTPException(
            status_code=500,
            detail=f"Internal server error: {str(e)}"
        )

# Download Template
@router.get("/{template_id}/download", tags=["Templates"])
async def download_template(
    template_id: str,
    user: dict = Depends(get_current_user)
):
    """
    Download a specific template file by ID.
    """
    try:
        logger.info(f"Attempting to download template {template_id} for user {user['id']}")

        # Fetch template details to get file_path and template_name
        response = supabase.from_("templates").select("file_path, template_name").eq("id", template_id).single().execute()

        if not response.data:
            raise HTTPException(
                status_code=404,
                detail=f"Template not found: {template_id}"
            )

        file_path = response.data["file_path"]
        template_name = response.data["template_name"]

        # Download file from Supabase storage
        try:
            # Supabase storage download returns bytes directly
            file_bytes = supabase.storage.from_("legal-templates").download(file_path)
            logger.info(f"Successfully retrieved file bytes for {file_path}")
            
            # Determine content type (can be more sophisticated if needed, e.g., using mimetypes)
            # For .docx files, the standard MIME type is application/vnd.openxmlformats-officedocument.wordprocessingml.document
            content_type = "application/vnd.openxmlformats-officedocument.wordprocessingml.document" if template_name.endswith(".docx") else "application/octet-stream"

            headers = {
                "Content-Disposition": f"attachment; filename=\"{template_name}\""
            }
            return Response(content=file_bytes, media_type=content_type, headers=headers)

        except Exception as download_error:
            logger.error(f"Error downloading file from storage {file_path}: {str(download_error)}")
            raise HTTPException(
                status_code=500,
                detail=f"Error downloading template file: {str(download_error)}"
            )

    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Unexpected error in download_template: {str(e)}\n{traceback.format_exc()}")
        raise HTTPException(
            status_code=500,
            detail=f"Internal server error: {str(e)}"
        )

@router.post("/upload-multi", tags=["Templates"])
async def upload_multiple_templates(
    files: Optional[List[UploadFile]] = File(None),
    templates_data: Optional[str] = Form(None),
    user: dict = Depends(get_current_user),
    credentials: HTTPAuthorizationCredentials = Depends(security)
) -> Dict[str, Any]:
    """
    Upload multiple templates via files or create them from provided content.

    The `templates_data` should be a JSON string of a list of dictionaries,
    each containing 'state', 'document_type', 'template_name', and 'content'.

    Example file upload:
    ```
    files: [template1.docx, template2.docx]
    ```

    Example content submission:
    ```
    templates_data: '[
        {"state": "California", "document_type": "Non-Disclosure Agreement", "template_name": "NDA_Content_1", "content": "Content for NDA 1"},
        {"state": "Texas", "document_type": "Employment Contract", "template_name": "Employment_Content_1", "content": "Content for Employment 1"}
    ]'
    ```
    """
    uploaded_templates_info = []

    if files:
        for file in files:
            if not file.filename.endswith('.docx'):
                logger.warning(f"Skipping {file.filename}: Only .docx files are supported for upload.")
                uploaded_templates_info.append({"filename": file.filename, "status": "skipped", "reason": "Only .docx files are supported"})
                continue

            try:
                # Extract content from DOCX
                content = ""
                file_content = await file.read()
                doc = Document(io.BytesIO(file_content))
                content = "\n".join([para.text for para in doc.paragraphs])
                logger.info(f"Extracted content from {file.filename}")

                # Determine state and document type from filename or metadata if available
                # For simplicity, let's assume state and document type are part of the filename or default
                # You might need a more sophisticated way to infer this or require it as part of form data
                # For now, we'll use a placeholder or extract from filename if possible
                state_name = "Unknown State"
                doc_type_name = "Unknown Document Type"
                
                # Attempt to extract state and document type from filename (basic example)
                name_parts = file.filename.replace('.docx', '').split('-')
                if len(name_parts) > 1:
                    state_name = name_parts[0]
                    doc_type_name = "-".join(name_parts[1:])

                state_response = supabase.from_("states").select("state_id").eq("state_name", state_name).single().execute()
                if not state_response.data:
                    logger.warning(f"State '{state_name}' not found for {file.filename}. Skipping.")
                    uploaded_templates_info.append({"filename": file.filename, "status": "skipped", "reason": f"Invalid state: {state_name}"})
                    continue
                state_id = state_response.data["state_id"]

                doc_type_response = supabase.from_("document_types").select("document_type_id").ilike("document_type_name", doc_type_name).single().execute()
                if not doc_type_response.data:
                    logger.warning(f"Document type '{doc_type_name}' not found for {file.filename}. Skipping.")
                    uploaded_templates_info.append({"filename": file.filename, "status": "skipped", "reason": f"Invalid document type: {doc_type_name}"})
                    continue
                doc_type_id = doc_type_response.data["document_type_id"]

                storage_path = f"legal-templates/{state_name}/{doc_type_name}/{file.filename}"
                
                # Upload to Supabase storage
                supabase.storage.from_("legal-templates").upload(storage_path, file_content, {"content-type": "application/vnd.openxmlformats-officedocument.wordprocessingml.document"})
                logger.info(f"File uploaded to storage: {storage_path}")

                # Insert into database
                template_data = {
                    "state_id": state_id,
                    "document_type_id": doc_type_id,
                    "template_name": file.filename,
                    "file_path": storage_path,
                    "content": content
                }
                response = supabase.from_("templates").insert(template_data).execute()
                uploaded_templates_info.append({"filename": file.filename, "status": "success", "template_id": response.data[0]['id'], "file_path": storage_path})
                logger.info(f"Template registered in database: {response.data[0]['id']}")

            except Exception as e:
                logger.error(f"Error processing file {file.filename}: {str(e)}\n{traceback.format_exc()}")
                uploaded_templates_info.append({"filename": file.filename, "status": "failed", "reason": str(e)})

    if templates_data:
        try:
            templates_to_create = json.loads(templates_data)
            for template_item in templates_to_create:
                state = template_item.get("state")
                document_type = template_item.get("document_type")
                template_name = template_item.get("template_name")
                content = template_item.get("content")

                if not all([state, document_type, template_name, content]):
                    logger.warning(f"Skipping template item due to missing data: {template_item}")
                    uploaded_templates_info.append({"template_item": template_item, "status": "skipped", "reason": "Missing required fields"})
                    continue
                
                try:
                    state_response = supabase.from_("states").select("state_id").eq("state_name", state).single().execute()
                    if not state_response.data:
                        logger.warning(f"State '{state}' not found for template '{template_name}'. Skipping.")
                        uploaded_templates_info.append({"template_name": template_name, "status": "skipped", "reason": f"Invalid state: {state}"})
                        continue
                    state_id = state_response.data["state_id"]

                    doc_type_response = supabase.from_("document_types").select("document_type_id").ilike("document_type_name", document_type).single().execute()
                    if not doc_type_response.data:
                        logger.warning(f"Document type '{document_type}' not found for template '{template_name}'. Skipping.")
                        uploaded_templates_info.append({"template_name": template_name, "status": "skipped", "reason": f"Invalid document type: {document_type}"})
                        continue
                    doc_type_id = doc_type_response.data["document_type_id"]

                    # For directly provided content, no file upload to storage, just DB entry
                    template_db_data = {
                        "state_id": state_id,
                        "document_type_id": doc_type_id,
                        "template_name": template_name,
                        "file_path": None, # No file path for directly provided content
                        "content": content
                    }
                    response = supabase.from_("templates").insert(template_db_data).execute()
                    uploaded_templates_info.append({"template_name": template_name, "status": "success", "template_id": response.data[0]['id']})
                    logger.info(f"Template (content-based) registered in database: {response.data[0]['id']}")

                except Exception as e:
                    logger.error(f"Error processing content-based template {template_name}: {str(e)}\n{traceback.format_exc()}")
                    uploaded_templates_info.append({"template_name": template_name, "status": "failed", "reason": str(e)})

        except json.JSONDecodeError:
            raise HTTPException(status_code=400, detail="Invalid JSON format for templates_data.")
        except Exception as e:
            logger.error(f"Unexpected error processing templates_data: {str(e)}\n{traceback.format_exc()}")
            raise HTTPException(status_code=500, detail=f"Internal server error processing templates_data: {str(e)}")

    if not files and not templates_data:
        raise HTTPException(status_code=400, detail="No files or template data provided.")

    return {"message": "Template upload/creation process completed.", "results": uploaded_templates_info}
