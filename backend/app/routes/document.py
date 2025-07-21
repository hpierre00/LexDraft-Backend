from fastapi import APIRouter, HTTPException, Depends, Query, File, UploadFile, Body, Form
from fastapi.security import HTTPAuthorizationCredentials
from app.models.document import DocumentCreate, DocumentUpdate, DocumentResponse
from app.models.schemas import DocumentGenerateRequest, ProfileInfo, ClientProfileResponse, ClientFolder, DocumentType, AreaOfLaw
from app.models.database import supabase
from app.utils.auth_utils import get_current_user, security
import logging
from typing import Dict, Any, List, Optional
import traceback
from fastapi.responses import FileResponse
import os
import tempfile
import markdown
from reportlab.lib.pagesizes import letter
from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer, ListFlowable, Image, PageBreak, Frame, flowables, HRFlowable, Table, TableStyle
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle, ListStyle
from reportlab.lib.enums import TA_CENTER
from reportlab.lib import colors
from starlette.background import BackgroundTask
from app.services.ai_agent_evaluate import evaluate_legal_document
from app.services.ai_agent_generate import generate_legal_document
from app.services.ai_agent_compliance import check_document_compliance
from app.services.ai_agent_enhance import enhance_document_with_ai
from app.utils.db_utils import get_profile, get_client_profile
from bs4 import BeautifulSoup
from docx import Document as DocxDocument
from PyPDF2 import PdfReader
from datetime import datetime
import json
from starlette.status import HTTP_201_CREATED

logger = logging.getLogger(__name__)
router = APIRouter()
from app.utils.db_utils import get_profile, get_client_profile # Import the utility functions to get profile data
from reportlab.lib import colors
from bs4 import BeautifulSoup
from docx import Document as DocxDocument # Import python-docx
from PyPDF2 import PdfReader # Import PyPDF2
from datetime import datetime
import json

# Configure logging
logger = logging.getLogger(__name__)

router = APIRouter()

# Utility function to save or update a document in Supabase
def save_document_to_supabase(user_id: str, title: str, content: str, status: str = "active", token: str = None) -> dict:
    """
    Save or update a document in the Supabase `documents` table.

    Parameters:
    - user_id (str): ID of the user.
    - title (str): Title of the document.
    - content (str): Content of the document.
    - status (str): Status of the document (default: "active").
    - token (str): User's session token for authentication.

    Returns:
    - dict: The saved document data.
    """
    try:
        # Set the auth token for this request
        if token:
            # For now, we'll use the same token for both access and refresh
            # In a production environment, you should get both tokens from the session
            supabase.auth.set_session(token, token)

        response = supabase.table("documents").insert({
            "user_id": user_id,
            "title": title,
            "content": content,
            "status": status
        }).execute()

        if not response.data:
            raise ValueError("No data returned from Supabase")

        return response.data[0]
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Error saving document: {str(e)}")

# Create Document
@router.post("/create", tags=["Documents"], response_model=DocumentResponse)
async def create_document(
    document_request: DocumentGenerateRequest = Body(...),
    user: dict = Depends(get_current_user),
    credentials: HTTPAuthorizationCredentials = Depends(security)
) -> Dict[str, Any]:
    """
    Create a new document using structured inputs and AI generation.
    """
    try:
        logger.info(f"Initiating document generation for user {user['id']} with title: {document_request.title}")
        
        user_profile_data = await get_profile(user["id"])
        if not user_profile_data:
            raise HTTPException(status_code=404, detail="User profile not found.")

        client_profile_data = None
        if document_request.client_profile_id:
            # Ensure the requesting user is an attorney to fetch client profiles
            if user_profile_data.get("role") != "attorney":
                raise HTTPException(status_code=403, detail="Only attorneys can generate documents for clients.")

            client_profile_data = await get_client_profile(
                attorney_id=user["id"],
                client_profile_id=str(document_request.client_profile_id)
            )
            if not client_profile_data:
                raise HTTPException(status_code=404, detail="Client profile not found or not accessible by this attorney.")

        # Generate document content using the AI agent
        generated_content = await generate_legal_document(
            notes=document_request.notes,
            user_id=user['id'],
            title=document_request.title,
            document_type=document_request.document_type,
            area_of_law=document_request.area_of_law,
            user_profile_data=user_profile_data, # Pass user profile data
            client_profile_id=document_request.client_profile_id,
            client_profile_data=client_profile_data, # Pass client profile data
            jurisdiction=document_request.jurisdiction,
            county=document_request.county,
            date_of_application=document_request.date_of_application,
            case_number=document_request.case_number,
        )
        logger.info(f"AI document generation complete for title: {document_request.title}")

        try:
            document_data = {
                "user_id": user["id"],
                "title": document_request.title,
                "content": generated_content,
                "status": "draft",
                "client_profile_id": str(document_request.client_profile_id) if document_request.client_profile_id else None
            }
            response = supabase.from_("documents").insert(document_data).execute()
            created_document = response.data[0]
            logger.info(f"Document created: {created_document['id']}")

            # Run compliance check after document creation
            compliance_result = await check_document_compliance(
                document_content=generated_content,
                jurisdiction=document_request.jurisdiction,
                document_type=document_request.document_type.value # Pass the enum value as string
            )
            
            # Update the document with compliance check results
            supabase.from_("documents").update({"compliance_check_results": compliance_result.dict()}).eq("id", created_document["id"]).execute()
            created_document["compliance_check_results"] = compliance_result.dict()
            logger.info(f"Compliance check completed and saved for document: {created_document['id']}")

            # Return the created document data
            return created_document

        except Exception as db_error:
            logger.error(f"Database insert/update failed: {str(db_error)}")
            raise HTTPException(
                status_code=500,
                detail=f"Error creating/evaluating document: {str(db_error)}"
            )
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Unexpected error in create_document: {str(e)}\n{traceback.format_exc()}")
        raise HTTPException(
            status_code=500,
            detail=f"Internal server error: {str(e)}"
        )

# Get All Documents
@router.get("/list", tags=["Documents"], response_model=list[DocumentResponse])
async def list_documents(
    status: Optional[str] = Query(None),
    search: Optional[str] = Query(None),
    user: dict = Depends(get_current_user)
) -> List[Dict[str, Any]]:
    """
    List documents with optional filters.
    """
    try:
        logger.info(f"Listing documents for user {user['id']} with filters: status={status}, search={search}")
        
        query = supabase.from_("documents").select("*").eq("user_id", user["id"])
        
        if status:
            query = query.eq("status", status)
        if search:
            query = query.or_(f"title.ilike.%{search}%,content.ilike.%{search}%")
            
        try:
            response = query.execute()
            logger.info(f"Found {len(response.data)} documents")
            
            # Post-process documents to ensure evaluation_response is a valid JSONB or None
            processed_documents = []
            for doc in response.data:
                if isinstance(doc.get("evaluation_response"), str):
                    try:
                        # Attempt to parse the string as JSON
                        doc["evaluation_response"] = json.loads(doc["evaluation_response"])
                    except json.JSONDecodeError:
                        # If it's not valid JSON, set to None
                        doc["evaluation_response"] = None
                processed_documents.append(doc)
            
            return processed_documents
        except Exception as query_error:
            logger.error(f"Query failed: {str(query_error)}")
            raise HTTPException(
                status_code=500,
                detail=f"Error fetching documents: {str(query_error)}"
            )
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Unexpected error in list_documents: {str(e)}\n{traceback.format_exc()}")
        raise HTTPException(
            status_code=500,
            detail=f"Internal server error: {str(e)}"
        )

# Get Single Document
@router.get("/{document_id}", tags=["Documents"], response_model=DocumentResponse)
async def get_document(
    document_id: str,
    user: dict = Depends(get_current_user)
) -> Dict[str, Any]:
    """
    Get a specific document by ID.
    """
    try:
        logger.info(f"Getting document {document_id} for user {user['id']}")
        
        try:
            # First try to get the document if user owns it
            response = supabase.from_("documents").select("*").eq("id", document_id).eq("user_id", user["id"]).execute()
            
            # If not found, check if user has access through team sharing or individual collaboration
            if not response.data:
                # Check if document exists first
                doc_exists = supabase.from_("documents").select("id, user_id, title, content, status, created_at, updated_at, evaluation_response").eq("id", document_id).single().execute()
                
                if not doc_exists.data:
                    raise HTTPException(
                        status_code=404,
                        detail=f"Document not found: {document_id}"
                    )
                
                # Check team access - user must be a team member and document must be shared with that team
                # First, get teams that have access to this document
                teams_with_access = supabase.from_("team_documents").select("team_id").eq("document_id", document_id).execute()
                
                team_access = None
                if teams_with_access.data:
                    # Check if user is a member of any of these teams
                    team_ids = [str(team["team_id"]) for team in teams_with_access.data]
                    if team_ids:
                        team_access = supabase.from_("team_members").select("id, team_id, role").eq("user_id", user["id"]).in_("team_id", team_ids).execute()
                
                # Check individual collaboration access
                collab_access = supabase.from_("document_collaborators").select("id, role").eq("document_id", document_id).eq("user_id", user["id"]).execute()
                
                has_team_access = team_access and team_access.data
                has_collab_access = collab_access.data
                
                if not has_team_access and not has_collab_access:
                    raise HTTPException(
                        status_code=403,
                        detail=f"Access denied: You don't have permission to view this document"
                    )
                
                # User has access, use the document data
                logger.info(f"Found shared document: {document_id}")
                return doc_exists.data
            else:
                logger.info(f"Found owned document: {document_id}")
                return response.data[0] if response.data else None
        except Exception as query_error:
            logger.error(f"Query failed: {str(query_error)}")
            raise HTTPException(
                status_code=500,
                detail=f"Error fetching document: {str(query_error)}"
            )
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Unexpected error in get_document: {str(e)}\n{traceback.format_exc()}")
        raise HTTPException(
            status_code=500,
            detail=f"Internal server error: {str(e)}"
        )

# Download Document
@router.get("/{document_id}/download", tags=["Documents"])
async def download_document(
    document_id: str,
    user: dict = Depends(get_current_user)
) -> FileResponse:
    """
    Download a specific document by ID as a PDF file.
    """
    temp_file_pdf_path = ""
    try:
        logger.info(f"Downloading document {document_id} for user {user['id']}")
        
        # First try to get the document if user owns it
        response = supabase.from_("documents").select("title", "content").eq("id", document_id).eq("user_id", user["id"]).execute()
        
        # If not found, check if user has access through team sharing or individual collaboration
        if not response.data:
            # Check if document exists first
            doc_exists = supabase.from_("documents").select("title", "content").eq("id", document_id).single().execute()
            
            if not doc_exists.data:
                raise HTTPException(
                    status_code=404,
                    detail=f"Document not found: {document_id}"
                )
            
            # Check team access
            # First, get teams that have access to this document
            teams_with_access = supabase.from_("team_documents").select("team_id").eq("document_id", document_id).execute()
            
            team_access = None
            if teams_with_access.data:
                # Check if user is a member of any of these teams
                team_ids = [str(team["team_id"]) for team in teams_with_access.data]
                if team_ids:
                    team_access = supabase.from_("team_members").select("id, team_id, role").eq("user_id", user["id"]).in_("team_id", team_ids).execute()
            
            # Check individual collaboration access
            collab_access = supabase.from_("document_collaborators").select("id, role").eq("document_id", document_id).eq("user_id", user["id"]).execute()
            
            has_team_access = team_access and team_access.data
            has_collab_access = collab_access.data
            
            if not has_team_access and not has_collab_access:
                logger.warning(f"User {user['id']} does not have access to document {document_id}")
                raise HTTPException(
                    status_code=403,
                    detail=f"Access denied: You don't have permission to download this document"
                )
            
            # User has access, use the document data
            document_data = doc_exists.data
        else:
            # User owns the document
            document_data = response.data[0] if response.data else None
            if not document_data:
                raise HTTPException(status_code=404, detail=f"Document not found: {document_id}")
        
        document_title = document_data["title"]
        document_content = document_data["content"]
        logger.info(f"Document {document_id} fetched successfully. Title: {document_title}")
        
        # Convert Markdown to HTML
        html_content = markdown.markdown(document_content)
        
        # Create a temporary PDF file
        with tempfile.NamedTemporaryFile(delete=False, suffix=".pdf") as temp_file_pdf:
            temp_file_pdf_path = temp_file_pdf.name
            doc = SimpleDocTemplate(
                temp_file_pdf_path,
                pagesize=letter,
                rightMargin=72,
                leftMargin=72,
                topMargin=72,
                bottomMargin=72
            )

            # Create and modify styles
            styles = getSampleStyleSheet()
            
            # Create new ParagraphStyle instances inheriting from default styles
            normal_style = ParagraphStyle(name='Normal', parent=styles['Normal'], fontName='Helvetica', fontSize=10, leading=12)
            h1_style = ParagraphStyle(name='h1', parent=styles['h1'], fontName='Helvetica-Bold', fontSize=20, leading=24, spaceAfter=12)
            h2_style = ParagraphStyle(name='h2', parent=styles['h2'], fontName='Helvetica-Bold', fontSize=16, leading=18, spaceAfter=9)
            h3_style = ParagraphStyle(name='h3', parent=styles['h3'], fontName='Helvetica-Bold', fontSize=14, leading=16, spaceAfter=7)

            # Add custom list item text style
            styles.add(ParagraphStyle(name='ListItemText', parent=normal_style, leftIndent=36, bulletIndent=12))

            # Define and add custom ListStyles
            styles.add(ListStyle(name='BulletList',
                                  parent=None,
                                  bulletFontName='Helvetica',
                                  bulletFontSize=10,
                                  bulletIndent=12,
                                  leftIndent=36,
                                  rightIndent=0,
                                  spaceBefore=6,
                                  spaceAfter=6,
                                  textColor=colors.black,
                                  bulletType='bullet'))

            styles.add(ListStyle(name='NumberList',
                                  parent=None,
                                  bulletFontName='Helvetica',
                                  bulletFontSize=10,
                                  bulletIndent=12,
                                  leftIndent=36,
                                  rightIndent=0,
                                  spaceBefore=6,
                                  spaceAfter=6,
                                  textColor=colors.black,
                                  bulletType='1'))

            # Parse the HTML content with BeautifulSoup to identify elements
            soup = BeautifulSoup(html_content, 'html.parser')

            story = []

            # Find all relevant elements (headings, paragraphs, lists, horizontal rules)
            # This approach avoids relying directly on soup.body.children which can be None
            content_elements = soup.find_all(['h1', 'h2', 'h3', 'p', 'ul', 'ol', 'hr'])

            if not content_elements and html_content.strip():
                # If no structured HTML elements are found, render the entire HTML content as a single paragraph
                story.append(Paragraph(html_content.strip(), normal_style))
            else:
                for element in content_elements:
                    # Handle headings
                    if element.name == 'h1':
                        story.append(Paragraph(element.get_text(), h1_style))
                    elif element.name == 'h2':
                        story.append(Paragraph(element.get_text(), h2_style))
                    elif element.name == 'h3':
                        story.append(Paragraph(element.get_text(), h3_style))
                    # Handle paragraphs
                    elif element.name == 'p':
                        story.append(Paragraph(element.get_text(), normal_style))
                    # Handle unordered and ordered lists
                    elif element.name in ['ul', 'ol']:
                        list_items = []
                        for li in element.find_all('li', recursive=False):
                            # Use styles['ListItemText'] as it's added via styles.add()
                            list_items.append(Paragraph(li.get_text(), styles['ListItemText']))

                        if element.name == 'ul':
                            story.append(ListFlowable(list_items,
                                                      bulletType='bullet',
                                                      start=None,
                                                      bulletDinkus='• ',
                                                      leftIndent=36,
                                                      bulletIndent=12,
                                                      spaceBefore=6,
                                                      spaceAfter=6,
                                                      style=styles['BulletList']))
                        else: # ol
                            story.append(ListFlowable(list_items,
                                                      bulletType='1',
                                                      start=None,
                                                      leftIndent=36,
                                                      bulletIndent=12,
                                                      spaceBefore=6,
                                                      spaceAfter=6,
                                                      style=styles['NumberList']))
                    # Handle horizontal rules
                    elif element.name == 'hr':
                        story.append(HRFlowable(width='100%', thickness=1, color=colors.black, spaceBefore=6, spaceAfter=6))
                    # No need for isinstance(element, str) here as find_all returns only Tag objects

            if not story and html_content.strip(): # Final fallback for plain content not caught above
                story.append(Paragraph(html_content.strip(), normal_style))

            # Build the PDF
            doc.build(story)
            logger.info(f"PDF generated successfully for document {document_id}")

        # Define a cleanup function to delete the temporary file after the response is sent
        def cleanup():
            os.remove(temp_file_pdf_path)
            logger.info(f"Temporary file {temp_file_pdf_path} cleaned up.")

        # Return the file as a response with background task for cleanup
        file_name = f"{document_title.replace(' ', '_')}.pdf"
        logger.info(f"Returning PDF file: {file_name}")
        return FileResponse(
            path=temp_file_pdf_path,
            filename=file_name,
            media_type="application/pdf",
            background=BackgroundTask(cleanup)
        )

    except HTTPException:
        if os.path.exists(temp_file_pdf_path):
            os.remove(temp_file_pdf_path)
        raise
    except Exception as e:
        logger.critical(f"Critical error in download_document: {str(e)}\n{traceback.format_exc()}")
        if os.path.exists(temp_file_pdf_path):
            os.remove(temp_file_pdf_path)
        raise HTTPException(
            status_code=500,
            detail=f"Internal server error: {str(e)}"
        )

@router.get("/{document_id}/download-docx", tags=["Documents"])
async def download_document_docx(
    document_id: str,
    user: dict = Depends(get_current_user)
) -> FileResponse:
    """
    Download a specific document by ID as a DOCX file.
    """
    temp_file_docx_path = ""
    try:
        logger.info(f"Downloading DOCX document {document_id} for user {user['id']}")
        
        # First try to get the document if user owns it
        response = supabase.from_("documents").select("title", "content").eq("id", document_id).eq("user_id", user["id"]).execute()
        
        # If not found, check if user has access through team sharing or individual collaboration
        if not response.data:
            # Check if document exists first
            doc_exists = supabase.from_("documents").select("title", "content").eq("id", document_id).single().execute()
            
            if not doc_exists.data:
                raise HTTPException(
                    status_code=404,
                    detail=f"Document not found: {document_id}"
                )
            
            # Check team access
            # First, get teams that have access to this document
            teams_with_access = supabase.from_("team_documents").select("team_id").eq("document_id", document_id).execute()
            
            team_access = None
            if teams_with_access.data:
                # Check if user is a member of any of these teams
                team_ids = [str(team["team_id"]) for team in teams_with_access.data]
                if team_ids:
                    team_access = supabase.from_("team_members").select("id, team_id, role").eq("user_id", user["id"]).in_("team_id", team_ids).execute()
            
            # Check individual collaboration access
            collab_access = supabase.from_("document_collaborators").select("id, role").eq("document_id", document_id).eq("user_id", user["id"]).execute()
            
            has_team_access = team_access and team_access.data
            has_collab_access = collab_access.data
            
            if not has_team_access and not has_collab_access:
                logger.warning(f"User {user['id']} does not have access to document {document_id}")
                raise HTTPException(
                    status_code=403,
                    detail=f"Access denied: You don't have permission to download this document"
                )
            
            # User has access, use the document data
            document_data = doc_exists.data
        else:
            # User owns the document
            document_data = response.data[0] if response.data else None
            if not document_data:
                raise HTTPException(status_code=404, detail=f"Document not found: {document_id}")
        
        document_title = document_data["title"]
        document_content = document_data["content"]
        
        # Create a new DOCX document
        doc = DocxDocument()
        
        # Convert Markdown to HTML
        html_content = markdown.markdown(document_content)
        soup = BeautifulSoup(html_content, 'html.parser')
        
        # Find all relevant elements (headings, paragraphs, lists, horizontal rules)
        content_elements = soup.find_all(['h1', 'h2', 'h3', 'p', 'ul', 'ol', 'hr'])

        if not content_elements and html_content.strip():
            # If no structured HTML elements are found, treat the entire content as a single paragraph
            doc.add_paragraph(html_content.strip())
        else:
            for element in content_elements:
                if element.name and element.name.startswith('h') and element.name[1:].isdigit():
                    level = int(element.name[1])
                    doc.add_heading(element.get_text(), level=min(level, 9)) # Max heading level in docx is 9
                elif element.name == 'p':
                    doc.add_paragraph(element.get_text())
                elif element.name == 'hr':
                    doc.add_paragraph('----------------------------------------------------------------------------------------------------------------')
                elif element.name in ['ul', 'ol']:
                    for li in element.find_all('li', recursive=False):
                        if element.name == 'ul':
                            doc.add_paragraph('• ' + li.get_text(), style='List Bullet')
                        else:
                            # For ordered lists, python-docx handles numbering with 'List Number' style
                            doc.add_paragraph(li.get_text(), style='List Number')
                # No need for isinstance(element, str) here as find_all returns only Tag objects
        
        # Create a temporary DOCX file
        with tempfile.NamedTemporaryFile(delete=False, suffix=".docx") as temp_file_docx:
            temp_file_docx_path = temp_file_docx.name
            doc.save(temp_file_docx_path)

        # Define a cleanup function to delete the temporary file after the response is sent
        def cleanup():
            os.remove(temp_file_docx_path)
            logger.info(f"Temporary DOCX file {temp_file_docx_path} cleaned up.")

        file_name = f"{document_title.replace(' ', '_')}.docx"
        return FileResponse(
            path=temp_file_docx_path,
            filename=file_name,
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            background=BackgroundTask(cleanup)
        )

    except HTTPException:
        if os.path.exists(temp_file_docx_path):
            os.remove(temp_file_docx_path)
        raise
    except Exception as e:
        logger.critical(f"Critical error in download_document_docx: {str(e)}\n{traceback.format_exc()}")
        if os.path.exists(temp_file_docx_path):
            os.remove(temp_file_docx_path)
        raise HTTPException(
            status_code=500,
            detail=f"Internal server error: {str(e)}"
        )

# Upload Evaluation
@router.post("/upload", tags=["Documents"], response_model=DocumentResponse)
async def upload_document_for_evaluation(
    file: UploadFile = File(...),
    user: dict = Depends(get_current_user)
) -> Dict[str, Any]:
    """
    Upload a document (PDF or DOCX) for content extraction and evaluation.
    """
    temp_file_path = ""
    try:
        logger.info(f"Received upload request for file: {file.filename}")
        
        # Save the uploaded file temporarily
        with tempfile.NamedTemporaryFile(delete=False, suffix=file.filename.split(".")[-1]) as temp_file:
            temp_file.write(await file.read())
            temp_file_path = temp_file.name
        logger.info(f"File saved temporarily to: {temp_file_path}")

        extracted_content = ""
        file_extension = file.filename.split(".")[-1].lower()

        if file_extension == "pdf":
            logger.info(f"Processing PDF file: {file.filename}")
            try:
                reader = PdfReader(temp_file_path)
                for page in reader.pages:
                    extracted_content += page.extract_text() or ""
                logger.info("PDF content extracted successfully.")
            except Exception as e:
                logger.error(f"Error extracting text from PDF: {str(e)}")
                raise HTTPException(status_code=500, detail=f"Error processing PDF file: {str(e)}")
        elif file_extension == "docx":
            logger.info(f"Processing DOCX file: {file.filename}")
            try:
                doc = DocxDocument(temp_file_path)
                for paragraph in doc.paragraphs:
                    extracted_content += paragraph.text + "\n"
                logger.info("DOCX content extracted successfully.")
            except Exception as e:
                logger.error(f"Error extracting text from DOCX: {str(e)}")
                raise HTTPException(status_code=500, detail=f"Error processing DOCX file: {str(e)}")
        else:
            logger.warning(f"Unsupported file type uploaded: {file.filename}")
            raise HTTPException(
                status_code=400,
                detail="Unsupported file type. Please upload a PDF or DOCX file."
            )
        
        if not extracted_content.strip():
            raise HTTPException(status_code=400, detail="No content extracted from the document.")

        # Perform AI evaluation
        logger.info(f"Performing AI evaluation for uploaded document: {file.filename}")
        evaluation_response = await evaluate_legal_document(document_content=extracted_content)
        logger.info(f"AI evaluation complete for uploaded document: {file.filename}")

        # Prepare document data for Supabase insertion
        document_data = {
            "user_id": user["id"],
            "title": file.filename,
            "content": extracted_content,
            "status": "evaluated",
            "evaluation_response": evaluation_response.dict() # Store the evaluation response as JSON
        }

        try:
            response = supabase.from_("documents").insert(document_data).execute()
            created_document = response.data[0]
            logger.info(f"Uploaded document saved to database: {created_document['id']}")

            # Run compliance check after evaluation
            compliance_result = await check_document_compliance(
                document_content=extracted_content,
                # For uploaded documents, jurisdiction and type might need to be inferred or provided separately
                # For now, we'll leave them as None or add logic to extract them if possible.
                jurisdiction=None,
                document_type=None
            )

            # Update the document with compliance check results
            supabase.from_("documents").update({"compliance_check_results": compliance_result.dict()}).eq("id", created_document["id"]).execute()
            created_document["compliance_check_results"] = compliance_result.dict()
            logger.info(f"Compliance check completed and saved for uploaded document: {created_document['id']}")
            
            # Re-fetch the document to ensure all latest fields, including evaluation_response and compliance_check_results, are present
            updated_document_response = supabase.from_("documents").select("*").eq("id", created_document["id"]).single().execute()
            if not updated_document_response.data:
                raise HTTPException(status_code=500, detail="Failed to retrieve updated document after compliance check.")

            # Removed debug logging
            return updated_document_response.data
        except Exception as db_error:
            logger.error(f"Database insert/update failed for uploaded document: {str(db_error)}")
            raise HTTPException(
                status_code=500,
                detail=f"Error saving/evaluating uploaded document: {str(db_error)}"
            )
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Unexpected error in upload_document_for_evaluation: {str(e)}\n{traceback.format_exc()}")
        raise HTTPException(
            status_code=500,
            detail=f"Internal server error: {str(e)}"
        )
    finally:
        if os.path.exists(temp_file_path):
            os.remove(temp_file_path)
            logger.info(f"Cleaned up temporary file: {temp_file_path}")

# Update Document
@router.put("/{document_id}", tags=["Documents"], response_model=DocumentResponse)
async def update_document(
    document_id: str,
    document_update: DocumentUpdate = Body(...),
    user: dict = Depends(get_current_user)
) -> Dict[str, Any]:
    """
    Update a document by ID.
    """
    try:
        logger.info(f"Updating document {document_id} for user {user['id']}")
        
        try:
            # First check if user can edit this document
            # Check if user owns the document
            doc_check = supabase.from_("documents").select("id, user_id").eq("id", document_id).single().execute()
            
            if not doc_check.data:
                raise HTTPException(status_code=404, detail=f"Document not found: {document_id}")
            
            can_edit = False
            
            # Owner can always edit
            if doc_check.data["user_id"] == user["id"]:
                can_edit = True
            else:
                # Check team collaboration permissions (editor+ role)
                # First, get teams that have access to this document
                teams_with_access = supabase.from_("team_documents").select("team_id").eq("document_id", document_id).execute()
                
                if teams_with_access.data:
                    # Check if user is a member of any of these teams with editor+ role
                    team_ids = [str(team["team_id"]) for team in teams_with_access.data]
                    if team_ids:
                        team_access = supabase.from_("team_members").select("id, team_id, role").eq("user_id", user["id"]).in_("team_id", team_ids).execute()
                        
                        if team_access.data:
                            # Check if user has editor+ role in any team with access to this document
                            for member in team_access.data:
                                user_role = member.get("role")
                                if user_role in ["owner", "admin", "editor"]:
                                    can_edit = True
                                    break
                
                # Check individual collaboration permissions (editor+ role)
                if not can_edit:
                    collab_access = supabase.from_("document_collaborators").select("role").eq("document_id", document_id).eq("user_id", user["id"]).execute()
                    
                    if collab_access.data:
                        user_role = collab_access.data[0]["role"]
                        if user_role in ["admin", "editor"]:
                            can_edit = True
            
            if not can_edit:
                raise HTTPException(
                    status_code=403,
                    detail="Access denied: You don't have permission to edit this document"
                )
            
            # Update only fields that are provided
            update_data = document_update.dict(exclude_unset=True)
            
            if not update_data:
                raise HTTPException(status_code=400, detail="No update data provided.")

            # Ensure updated_at is set
            update_data["updated_at"] = datetime.utcnow().isoformat() # Always update timestamp

            # Update the document (no user_id filter since we've already checked permissions)
            response = supabase.from_("documents").update(update_data).eq("id", document_id).execute()
            
            if not response.data:
                raise HTTPException(
                    status_code=500,
                    detail=f"Failed to update document: {document_id}"
                )

            logger.info(f"Document {document_id} updated successfully.")
            return response.data[0]
        except Exception as db_error:
            logger.error(f"Database update failed: {str(db_error)}")
            raise HTTPException(
                status_code=500,
                detail=f"Error updating document: {str(db_error)}"
            )
    except HTTPException:
        raise
    except Exception as e:
        logger.critical(f"Critical error in update_document for {document_id}: {str(e)}\n{traceback.format_exc()}")
        raise HTTPException(
            status_code=500,
            detail=f"Internal server error: {str(e)}"
        )

# Delete (Archive) Document
@router.delete("/{document_id}", tags=["Documents"])
async def delete_document(
    document_id: str,
    user: dict = Depends(get_current_user)
) -> Dict[str, str]:
    """
    Delete a document by ID.
    """
    try:
        logger.info(f"Deleting document {document_id} for user {user['id']}")
        
        try:
            response = supabase.from_("documents").delete().eq("id", document_id).eq("user_id", user["id"]).execute() # Filter by user_id
            
            if not response.data:
                raise HTTPException(
                    status_code=404,
                    detail=f"Document not found: {document_id}"
                )
            logger.info(f"Document deleted: {document_id}")
            return {"message": "Document deleted successfully"}
        except Exception as delete_error:
            logger.error(f"Delete failed: {str(delete_error)}")
            raise HTTPException(
                status_code=500,
                detail=f"Error deleting document: {str(delete_error)}"
            )
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Unexpected error in delete_document: {str(e)}\n{traceback.format_exc()}")
        raise HTTPException(
            status_code=500,
            detail=f"Internal server error: {str(e)}"
        )

@router.put("/{document_id}/archive", tags=["Documents"])
async def archive_document(
    document_id: str,
    user: dict = Depends(get_current_user)
) -> Dict[str, Any]:
    """
    Archives a specific document.
    """
    try:
        logger.info(f"Archiving document {document_id} for user {user['id']}")
        
        try:
            # Set document status to 'archived'
            response = supabase.from_("documents").update({"status": "archived"}).eq("id", document_id).eq("user_id", user["id"]).execute() # Filter by user_id
            
            if not response.data:
                raise HTTPException(
                    status_code=404,
                    detail=f"Document not found: {document_id}"
                )
            logger.info(f"Document archived: {document_id}")
            return response.data[0]
        except Exception as archive_error:
            logger.error(f"Archive failed: {str(archive_error)}")
            raise HTTPException(
                status_code=500,
                detail=f"Error archiving document: {str(archive_error)}"
            )
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Unexpected error in archive_document: {str(e)}\n{traceback.format_exc()}")
        raise HTTPException(
            status_code=500,
            detail=f"Internal server error: {str(e)}"
        )

# New endpoint to manually run compliance check
@router.post("/{document_id}/run-compliance", tags=["Documents"], response_model=DocumentResponse)
async def run_compliance_check_on_document(
    document_id: str,
    user: dict = Depends(get_current_user)
) -> Dict[str, Any]:
    """
    Manually runs the compliance check on an existing document.
    """
    try:
        logger.info(f"Running compliance check on document {document_id} for user {user['id']}")

        # Fetch the document content
        response = supabase.from_("documents").select("content", "title", "created_at", "client_profile_id").eq("id", document_id).eq("user_id", user["id"]).single().execute()
        if not response.data:
            raise HTTPException(status_code=404, detail=f"Document not found or not accessible: {document_id}")

        document_content = response.data["content"]
        # For re-running, we might not have document_type or jurisdiction readily available
        # You may want to store these with the document initially, or infer them here if possible.
        # For now, we'll pass None for simplicity if not available from document metadata.
        document_title = response.data["title"]
        created_at = response.data["created_at"]
        client_profile_id = response.data["client_profile_id"]

        # Run compliance check
        compliance_result = await check_document_compliance(
            document_content=document_content,
            jurisdiction=None, # To be improved if jurisdiction is stored with the document
            document_type=None # To be improved if document_type is stored with the document
        )

        # Update the document with new compliance results
        update_response = supabase.from_("documents").update({"compliance_check_results": compliance_result.dict()}).eq("id", document_id).execute()
        if not update_response.data:
            raise HTTPException(status_code=500, detail="Failed to update document with compliance results.")

        # Fetch the updated document to return the full response model
        updated_document = supabase.from_("documents").select("*").eq("id", document_id).eq("user_id", user["id"]).single().execute()
        return updated_document.data

    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Unexpected error in run_compliance_check_on_document: {str(e)}\n{traceback.format_exc()}")
        raise HTTPException(
            status_code=500,
            detail=f"Internal server error: {str(e)}"
        )
    finally:
        pass
# Enhance with AI: Upload and Enhance a New Document
@router.post("/enhance/upload", tags=["Documents"], response_model=DocumentResponse, status_code=HTTP_201_CREATED)
async def enhance_upload_document(
    file: UploadFile = File(...),
    instructions: Optional[str] = Form(None),
    user: dict = Depends(get_current_user)
) -> Dict[str, Any]:
    """
    Upload a document (PDF, DOCX, TXT), enhance it with AI, and save as a new document.
    """
    temp_file_path = ""
    try:
        logger.info(f"Received enhance-upload request for file: {file.filename}")
        # Save the uploaded file temporarily
        with tempfile.NamedTemporaryFile(delete=False, suffix=file.filename.split(".")[-1]) as temp_file:
            temp_file.write(await file.read())
            temp_file_path = temp_file.name
        logger.info(f"File saved temporarily to: {temp_file_path}")

        extracted_content = ""
        file_extension = file.filename.split(".")[-1].lower()
        if file_extension == "pdf":
            reader = PdfReader(temp_file_path)
            for page in reader.pages:
                extracted_content += page.extract_text() or ""
        elif file_extension == "docx":
            doc = DocxDocument(temp_file_path)
            for paragraph in doc.paragraphs:
                extracted_content += paragraph.text + "\n"
        elif file_extension == "txt":
            with open(temp_file_path, "r", encoding="utf-8") as f:
                extracted_content = f.read()
        else:
            raise HTTPException(status_code=400, detail="Unsupported file type. Please upload a PDF, DOCX, or TXT file.")

        if not extracted_content.strip():
            raise HTTPException(status_code=400, detail="No content extracted from the document.")

        # Enhance with AI
        enhanced_content = enhance_document_with_ai(extracted_content, instructions)

        # Save to database
        document_data = {
            "user_id": user["id"],
            "title": file.filename,
            "content": enhanced_content,
            "status": "enhanced"
        }
        response = supabase.from_("documents").insert(document_data).execute()
        created_document = response.data[0]
        logger.info(f"Enhanced document saved to database: {created_document['id']}")
        return created_document
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Unexpected error in enhance_upload_document: {str(e)}\n{traceback.format_exc()}")
        raise HTTPException(status_code=500, detail=f"Internal server error: {str(e)}")
    finally:
        if temp_file_path and os.path.exists(temp_file_path):
            os.remove(temp_file_path)

# Enhance with AI: Enhance an Existing Document
@router.post("/enhance/{document_id}", tags=["Documents"], response_model=DocumentResponse)
async def enhance_existing_document(
    document_id: str,
    instructions: str = Body(..., embed=True),
    user: dict = Depends(get_current_user)
) -> Dict[str, Any]:
    """
    Enhance an existing document by providing instructions. Updates the document in place.
    """
    try:
        # Fetch the document
        doc_response = supabase.from_("documents").select("*").eq("id", document_id).single().execute()
        if not doc_response.data:
            raise HTTPException(status_code=404, detail="Document not found.")
        document = doc_response.data
        if document["user_id"] != user["id"]:
            raise HTTPException(status_code=403, detail="You do not have permission to enhance this document.")

        # Enhance with AI
        enhanced_content = enhance_document_with_ai(document["content"], instructions)

        # Update the document
        update_data = {"content": enhanced_content, "status": "enhanced"}
        supabase.from_("documents").update(update_data).eq("id", document_id).execute()
        document.update(update_data)
        logger.info(f"Document {document_id} enhanced and updated.")
        return document
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Unexpected error in enhance_existing_document: {str(e)}\n{traceback.format_exc()}")
        raise HTTPException(status_code=500, detail=f"Internal server error: {str(e)}")